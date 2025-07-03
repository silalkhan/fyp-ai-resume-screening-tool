import os
import logging
import re
import tempfile
import subprocess
from io import StringIO

# Configure logging
logger = logging.getLogger(__name__)

# Import potentially problematic modules inside functions to avoid initialization errors
def _load_pdf_module():
    try:
        from PyPDF2 import PdfReader
        return PdfReader
    except ImportError:
        logger.warning("PyPDF2 not available, PDF extraction may be limited")
        return None

def _load_docx_module():
    try:
        import docx
        return docx
    except ImportError:
        logger.warning("python-docx not available, DOCX extraction may be limited")
        return None

def extract_text(file_path):
    """
    Extract text from resume file (PDF or DOCX)
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text or None if extraction failed
    """
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
            
        file_extension = os.path.splitext(file_path)[1].lower()
        logger.info(f"Extracting text from {file_path} with extension {file_extension}")
        
        # Basic file validation
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.error(f"Empty file: {file_path}")
            return None
        elif file_size > 10 * 1024 * 1024:  # 10MB
            logger.warning(f"Very large file ({file_size} bytes): {file_path}")
        
        # Check if file is readable
        try:
            with open(file_path, 'rb') as f:
                first_bytes = f.read(8)
                logger.info(f"First bytes of file: {first_bytes.hex()}")
        except Exception as read_error:
            logger.error(f"Could not read file: {str(read_error)}")
            return None
            
        # PDF extraction
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        # DOCX extraction
        elif file_extension == '.docx':
            # Try multiple methods to extract text from DOCX
            result = extract_text_from_docx(file_path)
            
            # If the result is empty or very short, try a secondary method
            if not result or len(result.strip()) < 100:
                logger.warning(f"Primary DOCX extraction yielded limited text ({len(result) if result else 0} chars), trying alternative methods")
                try:
                    alt_result = extract_text_from_docx_alternative(file_path)
                    if alt_result and len(alt_result) > len(result or ""):
                        logger.info(f"Alternative DOCX extraction succeeded with {len(alt_result)} chars")
                        return alt_result
                except Exception as alt_error:
                    logger.error(f"Alternative DOCX extraction failed: {str(alt_error)}")
            
            return result
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def extract_text_from_pdf(file_path):
    """Extract text from PDF using multiple methods for robustness"""
    text = ""
    
    # First try PyPDF2
    try:
        logger.info(f"Extracting text from PDF using PyPDF2: {file_path}")
        PdfReader = _load_pdf_module()
        if not PdfReader:
            raise ImportError("PyPDF2 not available")
            
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        
        # If we got reasonable text, return it
        if len(text.strip()) > 100:
            logger.info(f"Successfully extracted {len(text)} characters with PyPDF2")
            return clean_text(text)
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {str(e)}")
    
    # If PyPDF2 failed or extracted too little text, try pdftotext if available
    try:
        logger.info(f"Trying pdftotext for {file_path}")
        with tempfile.NamedTemporaryFile(suffix='.txt') as temp_txt:
            # Try to use pdftotext command line tool if available
            result = subprocess.run(
                ['pdftotext', '-layout', file_path, temp_txt.name],
                capture_output=True,
                text=True,
                check=False
            )
            
            if result.returncode == 0:
                with open(temp_txt.name, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                
                if len(text.strip()) > 0:
                    logger.info(f"Successfully extracted {len(text)} characters with pdftotext")
                    return clean_text(text)
            else:
                logger.warning(f"pdftotext failed: {result.stderr}")
    except (FileNotFoundError, subprocess.SubprocessError) as e:
        logger.warning(f"pdftotext extraction failed: {str(e)}")
    
    # If we have some text from PyPDF2 but it's not ideal, return it anyway
    if len(text.strip()) > 0:
        logger.info(f"Returning partial text ({len(text)} chars) from PyPDF2")
        return clean_text(text)
    
    logger.error(f"Failed to extract text from PDF: {file_path}")
    return None

def extract_text_from_docx(file_path):
    """Extract text from DOCX file with multiple fallback methods"""
    text = ""
    
    # Method 1: python-docx library
    try:
        logger.info(f"Extracting text from DOCX using python-docx: {file_path}")
        docx = _load_docx_module()
        if not docx:
            raise ImportError("python-docx not available")
        
        # Log file details before opening
        file_size = os.path.getsize(file_path)
        logger.info(f"DOCX file size: {file_size} bytes")
        
        # Try opening document with detailed error reporting
        try:
            doc = docx.Document(file_path)
            logger.info(f"DOCX opened successfully. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
        except Exception as doc_error:
            logger.error(f"Error opening DOCX with python-docx: {str(doc_error)}")
            import traceback
            logger.error(traceback.format_exc())
            raise
            
        # Extract paragraphs with error tracking
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            try:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    if i < 5:  # Log first few paragraphs for debugging
                        logger.info(f"Paragraph {i}: {text[:50]}...")
            except Exception as para_error:
                logger.error(f"Error extracting paragraph {i}: {str(para_error)}")
        
        logger.info(f"Extracted {len(paragraphs)} non-empty paragraphs from DOCX")
        
        # Extract tables with error tracking
        tables_text = []
        for i, table in enumerate(doc.tables):
            try:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            except Exception as table_error:
                logger.error(f"Error extracting table {i}: {str(table_error)}")
        
        logger.info(f"Extracted {len(tables_text)} table rows from DOCX")
        
        # Combine text
        text = '\n'.join(paragraphs)
        if tables_text:
            text += '\n\n' + '\n'.join(tables_text)
            
        if len(text.strip()) > 100:
            logger.info(f"Successfully extracted {len(text)} characters from DOCX using python-docx")
            return clean_text(text)
    except Exception as e:
        logger.warning(f"python-docx extraction failed: {str(e)}")
    
    # Fallback Method 2: Try using textract if available
    if not text.strip() or len(text.strip()) < 100:
        try:
            import textract
            logger.info(f"Trying textract for DOCX extraction: {file_path}")
            extracted_text = textract.process(file_path).decode('utf-8', errors='ignore')
            if extracted_text and len(extracted_text.strip()) > 0:
                logger.info(f"Successfully extracted {len(extracted_text)} characters using textract")
                return clean_text(extracted_text)
        except ImportError:
            logger.warning("textract not available for DOCX extraction fallback")
        except Exception as e:
            logger.warning(f"textract extraction failed: {str(e)}")
    
    # Fallback Method 3: Try using docx2txt if available
    if not text.strip() or len(text.strip()) < 100:
        try:
            import docx2txt
            logger.info(f"Trying docx2txt for DOCX extraction: {file_path}")
            extracted_text = docx2txt.process(file_path)
            if extracted_text and len(extracted_text.strip()) > 0:
                logger.info(f"Successfully extracted {len(extracted_text)} characters using docx2txt")
                return clean_text(extracted_text)
        except ImportError:
            logger.warning("docx2txt not available for DOCX extraction fallback")
        except Exception as e:
            logger.warning(f"docx2txt extraction failed: {str(e)}")
    
    # If we got at least something from python-docx, return that
    if text.strip():
        logger.info(f"Returning partial text ({len(text)} chars) from python-docx")
        return clean_text(text)
    
    logger.error(f"All DOCX extraction methods failed for: {file_path}")
    return None

def extract_text_from_docx_alternative(file_path):
    """
    Alternative method to extract text from DOCX using direct XML parsing
    for cases where standard libraries fail
    """
    try:
        logger.info(f"Attempting alternative DOCX extraction for: {file_path}")
        import zipfile
        import xml.etree.ElementTree as ET
        
        # DOCX files are ZIP archives containing XML
        text_content = []
        
        # Check if file can be opened as a ZIP
        if not zipfile.is_zipfile(file_path):
            logger.error(f"File is not a valid ZIP/DOCX: {file_path}")
            return None
            
        # Extract document.xml which contains the main content
        with zipfile.ZipFile(file_path) as docx_zip:
            # List the contents for debugging
            file_list = docx_zip.namelist()
            logger.info(f"ZIP contents: {file_list[:10]}...")
            
            # Look for document.xml in standard location
            doc_xml_path = 'word/document.xml'
            if doc_xml_path not in file_list:
                logger.error(f"document.xml not found in DOCX")
                return None
                
            # Extract and parse XML
            with docx_zip.open(doc_xml_path) as doc_xml:
                tree = ET.parse(doc_xml)
                root = tree.getroot()
                
                # DOCX XML namespace
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Extract text from paragraphs
                for paragraph in root.findall('.//w:p', ns):
                    p_text = []
                    for text_element in paragraph.findall('.//w:t', ns):
                        if text_element.text:
                            p_text.append(text_element.text)
                    if p_text:
                        text_content.append(''.join(p_text))
        
        # Combine all paragraphs
        full_text = '\n'.join(text_content)
        logger.info(f"Alternative extraction found {len(text_content)} paragraphs, {len(full_text)} chars")
        
        if not full_text.strip():
            logger.warning("Alternative extraction returned empty text")
            return None
            
        return clean_text(full_text)
    except Exception as e:
        logger.error(f"Alternative DOCX extraction failed: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def clean_text(text):
    """Clean and normalize extracted text"""
    if not text:
        return ""
        
    # Remove non-printable characters
    text = ''.join(c for c in text if c.isprintable() or c in '\n\t')
    
    # Replace tabs with spaces
    text = text.replace('\t', ' ')
    
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove excessive whitespace within lines
    lines = text.split('\n')
    clean_lines = []
    for line in lines:
        clean_lines.append(' '.join(word for word in line.split() if word))
    
    # Rejoin with newlines
    text = '\n'.join(clean_lines)
    
    # Remove very long tokens (likely garbage)
    text = ' '.join(word for word in text.split() if len(word) < 30)
    
    return text.strip()

# Add a standalone test function for direct invocation
def test_extraction(file_path):
    """Test extraction on a specific file and print results"""
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return False
        
    print(f"Testing extraction on: {file_path}")
    extracted_text = extract_text(file_path)
    
    if extracted_text:
        print(f"Successfully extracted {len(extracted_text)} characters")
        print("\nPreview (first 500 chars):")
        print("-" * 80)
        print(extracted_text[:500] + ("..." if len(extracted_text) > 500 else ""))
        print("-" * 80)
        return True
    else:
        print("Extraction failed!")
        return False

# Allow running as standalone script for testing
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python extract_text.py <file_path>")
        sys.exit(1)
    
    success = test_extraction(sys.argv[1])
    sys.exit(0 if success else 1)